from django.shortcuts import render
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import Sum,Count,Case,When,Value,BooleanField,Max
from django.views import View
from django.http import JsonResponse, FileResponse, HttpResponse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils.crypto import get_random_string
from django.conf import settings
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync
from django.utils import timezone

from utils.mixins import CompanyOwnershipRequiredMixin

from common.utils.export_utils import BaseExporter
from common.utils.websocket_utils import send_alert
from common.models import ExportProcess
from leasing.models import Lease,Installment
from leasing.utils.lease_utils import get_future_payments
from contracts.models import TerminationWarningNotice, CommitteeForm,WarningNotice
from partners.models import Partner
from risk.utils.filter_utils import cancellation_warning_notice_exists

import os
import json
import pandas as pd
from collections import Counter
from decimal import Decimal
from datetime import datetime
from docxtpl import DocxTemplate
import io, zipfile
from copy import deepcopy
from docx.oxml.ns import qn

# Create your views here.

def format_currency(value):
    return "{:,.2f}".format(value).replace(",", "X").replace(".", ",").replace("X", ".")

class ExportToTerminatedRiskPartnersForSMSView(LoginRequiredMixin,View):
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)

        exporter = BaseExporter(
            user_id=request.user.id,
            app="risk",
            model_name="ToTerminatedRiskPartnerForSMS",
            file_name=f"{datetime.today().strftime('%d-%m-%Y')}-fesih-edilecekler-sms.xlsx",
            export_url="/risk/to_terminated_risk_partners_excel_for_sms",
            params={"project":data.get('project')}
        )

        send_alert({"message":"Excel dosyası hazırlanıyor...",'status':'success'},room=f"private_{request.user.id}")
            
        exporter.start_export()

        return HttpResponse(status=200)

class ToTerminatedRiskPartnersExcelForSMSView(LoginRequiredMixin,View):
    def get(self, request, *args, **kwargs):
        file_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{datetime.today().strftime('%d-%m-%Y')}-fesih-edilecekler-sms.xlsx")
      
        if not os.path.exists(file_path):
            return JsonResponse({'message': 'File not found!','status':'error'}, status=404)

        objs = ExportProcess.objects.filter(status = "in_progress")
        for obj in objs:
            obj.status = "completed"
            obj.save()

        return FileResponse(open(file_path, 'rb'))


class ExportToTerminatedRiskPartnersView(LoginRequiredMixin,View):
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)

        exporter = BaseExporter(
            user_id=request.user.id,
            app="risk",
            model_name="ToTerminatedRiskPartner",
            file_name=f"{datetime.today().strftime('%d-%m-%Y')}-fesih-edilecekler.xlsx",
            export_url="/risk/to_terminated_risk_partners_excel",
            params={"project":data.get('project')}
        )

        send_alert({"message":"Excel dosyası hazırlanıyor...",'status':'success'},room=f"private_{request.user.id}")
            
        exporter.start_export()

        return HttpResponse(status=200)

class ToTerminatedRiskPartnersExcelView(LoginRequiredMixin,View):
    def get(self, request, *args, **kwargs):
        file_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{datetime.today().strftime('%d-%m-%Y')}-fesih-edilecekler.xlsx")
      
        if not os.path.exists(file_path):
            return JsonResponse({'message': 'File not found!','status':'error'}, status=404)

        objs = ExportProcess.objects.filter(status = "in_progress")
        for obj in objs:
            obj.status = "completed"
            obj.save()

        return FileResponse(open(file_path, 'rb'))

class CreateTerminationWarningNoticeStatusView(LoginRequiredMixin,CompanyOwnershipRequiredMixin,View):
    model = Lease
    
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)

        lease = Lease.objects.select_related().filter(uuid = data.get('uuid')).first()

        if lease:
            # kapsamlı ihtar model işlemleri
            if not TerminationWarningNotice.objects.filter(contract = lease.contract).exists():
                obj = TerminationWarningNotice.objects.create(
                    company = lease.company,
                    contract = lease.contract,
                )
            else:
                obj = TerminationWarningNotice.objects.filter(contract = lease.contract).first()

            # word işlemleri
            file_name = lease.contract.code.replace("/","-")
            doc = DocxTemplate(f"files/fesih-ihtar.docx")

            def format_currency(value):
                    return "{:,.2f}".format(value).replace(",", "X").replace(".", ",").replace("X", ".")
            
            if lease.contract.partner.is_commercial:
                if lease.contract.partner.tc_vkn_no and len(lease.contract.partner.tc_vkn_no) > 0:
                    tc_vkn_no = lease.contract.partner.tc_vkn_no
                elif lease.contract.partner.vat_no and len(lease.contract.partner.vat_no) > 0:
                    tc_vkn_no = lease.contract.partner.vat_no
                else:
                    tc_vkn_no = ''
            else:
                tc_vkn_no = lease.contract.partner.tc_vkn_no if lease.contract.partner.tc_vkn_no else ''

            context = {
                "isim": lease.contract.partner.name,
                "tc_vkn_no": tc_vkn_no,
                "adres": lease.contract.partner.address,
                "sozlesme_tarih": lease.signature_date.strftime('%d.%m.%Y') if lease.signature_date else '',
                "sozlesme_no": lease.contract.code,
                "odenen_tutar": format_currency(obj.paid_amount),
                "kesinti_tutar": format_currency(obj.deduction_amount),
                "toplam_tutar": format_currency(obj.paid_amount - obj.deduction_amount),
            }
            doc.render(context)

            files_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{file_name}.docx")
            doc.save(files_path)

        if files_path:
            return FileResponse(open(files_path, 'rb'), as_attachment=True)
        
        return JsonResponse({'message': 'Dosya bulunamadı!','status':'error'}, status=400)

class GetTerminationWarningNoticeView(LoginRequiredMixin,View):
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)
        uuid = data.get('uuid')

        lease = Lease.objects.select_related().filter(uuid = uuid).first()

        file_name = lease.contract.code.replace("/","-")
        file_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{file_name}.docx")
        
        if not os.path.exists(file_path):
            return JsonResponse({'message': 'Dosya bulunamadı!','status':'error'}, status=404)
        
        return FileResponse(open(file_path, 'rb'))
    
class CreateCommitteeFormStatusView(LoginRequiredMixin,View):
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)
        
        partner = Partner.objects.select_related().filter(uuid = data.get('uuid')).first()
        leases = Lease.objects.select_related().filter(uuid__in = data.get('uuids', []))

        if partner:
            # kapsamlı ihtar model işlemleri
            if not CommitteeForm.objects.filter(partner = partner).exists():
                obj = CommitteeForm.objects.create(
                    company = partner.company,
                    partner = partner,
                )
            else:
                obj = CommitteeForm.objects.filter(partner = partner).first()

            # word işlemleri
            file_name = f"{partner.name.lower().replace(' ', '_')}-{str(partner.crm_code)}-komite-formu"
            doc = DocxTemplate(f"files/komite-formu.docx")

            #hesaplamalar
            risk_next_installments_total_amount = Decimal('0.00')
            risk_overdue_amount_total = Decimal('0.00')
            risk_total_risk_amount_total = Decimal('0.00')
            risk_paid_amount_total = Decimal('0.00')
            refund_installment_amount_total = Decimal('0.00')
            muhtelif_transactions_total_amount = Decimal('0.00')
            teblig_tarihi_list = []
            for lease in leases:
                next_installments = Installment.objects.select_related("lease").filter(lease = lease, type__in = ['1'], payment_date__gte = datetime.today()).only("amount").aggregate(total_amount=Sum('amount'))
                borc_muhtelif_transactions = lease.lease_trade_transactions.filter(amount_type = '1', posting_group_name='Muhtelif Masraf').aggregate(total_amount=Sum('amount'))
                alacak_muhtelif_transactions = lease.lease_trade_transactions.filter(amount_type = '0', posting_group_name='Muhtelif Masraf').aggregate(total_amount=Sum('amount'))
                muhtelif_transactions_total_amount += (borc_muhtelif_transactions['total_amount'] or Decimal('0.00')) - (alacak_muhtelif_transactions['total_amount'] or Decimal('0.00'))
                
                risk_next_installments_total_amount += next_installments['total_amount'] or Decimal('0.00')
                risk_overdue_amount_total += lease.overdue_amount
                risk_total_risk_amount_total += (next_installments['total_amount'] or Decimal('0.00')) + lease.overdue_amount
                risk_paid_amount_total += lease.paid_amount
                refund_installment_amount_total += lease.installment_amount
                warning_notice = WarningNotice.objects.filter(
                    contract=lease.contract,
                    service_date__isnull=False,
                    official_cancellation_date__lte=timezone.now().date(),
                    state__in=['Yeni', 'Geçerli']
                ).order_by('-created_date').first()
                teblig_tarihi_list.append(warning_notice.service_date if warning_notice else None)

            teblig_tarihi_counts = Counter(d.strftime('%d.%m.%Y') for d in teblig_tarihi_list if d is not None)
            teblig_tarihi_list = [{'tarih': tarih, 'adet': str(adet), 'tarih_eki': "'dır" if datetime.strptime(tarih, '%d.%m.%Y').year == 2026 else "'dir"} for tarih, adet in teblig_tarihi_counts.items()]

            tahsil_edilen = risk_paid_amount_total
            ihtar_masrafi = Decimal('3200.00') if not partner.kep else Decimal('1000.00')*Decimal(str(leases.count()))
            takip_masrafi = risk_overdue_amount_total*Decimal('0.1') if risk_overdue_amount_total > muhtelif_transactions_total_amount else muhtelif_transactions_total_amount
            vazgecme_akcesi = refund_installment_amount_total*Decimal('0.1')
            kalan = risk_paid_amount_total - ihtar_masrafi - takip_masrafi - vazgecme_akcesi

            context = {
                "isim": partner.name,
                "tarih": datetime.today().strftime('%d.%m.%Y'),
                "kalan": format_currency(kalan),
                "pb": leases[0].currency.code if leases and leases[0].currency.code != "TRY" else "TL",
                "teblig": "kep" if partner.kep else "noter",
                #"teblig_tarihi_list": "\n".join(teblig_tarihi['tarih_eki'] for teblig_tarihi in teblig_tarihi_list),
                "ihtar_kelimesi": "ihtarlari" if leases.count() > 1 else "ihtarı",
            }

            doc.render(context)

            files_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{file_name}.docx")
            doc.save(files_path)

            # tabloya satır ekleme
            from docx import Document
            document = Document(files_path)

            for i, paragraph in enumerate(document.paragraphs):
                if '<<teblig_tarihi_list>>' in paragraph.text:  # şablonda bu yer tutucuyu koy
                    for teblig in teblig_tarihi_list:
                        new_para = deepcopy(paragraph._p)
                        new_para.find('.//' + qn('w:t')).text = (
                            f"{teblig['adet']} dönem için tebliğ {"tarihleri" if leases.count() > 1 else "tarihi"} "
                            f"{teblig['tarih']}{teblig['tarih_eki']}"
                        )
                        paragraph._p.addprevious(new_para)
                    paragraph._p.getparent().remove(paragraph._p)
                    break

            project_table = document.tables[2]
            template_lease_row_project_table = project_table.rows[2]
            template_total_row_project_table = project_table.rows[3]

            risk_table = document.tables[3]
            template_lease_row_risk_table = risk_table.rows[2]
            template_total_row_risk_table = risk_table.rows[3]

            refund_table = document.tables[4]
            template_installment_amount_row_refund_table = refund_table.rows[2]
            template_paid_amount_row_refund_table = refund_table.rows[3]
            template_warning_notice_row_refund_table = refund_table.rows[4]
            template_tracking_row_refund_table = refund_table.rows[5]
            template_payout_row_refund_table = refund_table.rows[6]
            template_remaining_row_refund_table = refund_table.rows[7]

            for lease in leases:
                #projeler tablosu___________________________________________________________________________________________
                lease_row_project_table = deepcopy(template_lease_row_project_table._tr)
                cells_lease_project_table = lease_row_project_table.findall(qn('w:tc'))
                cells_lease_project_table[0].find('.//' + qn('w:t')).text = lease.contract.vendor.name if lease.contract and lease.contract.vendor else ""
                cells_lease_project_table[1].find('.//' + qn('w:t')).text = lease.item.stock_name if lease.item else ""
                cells_lease_project_table[2].find('.//' + qn('w:t')).text = lease.contract.code if lease.contract else ""
                cells_lease_project_table[3].find('.//' + qn('w:t')).text = lease.get_lease_status_display()
                cells_lease_project_table[4].find('.//' + qn('w:t')).text = lease.signature_date.strftime('%d.%m.%Y') if lease.signature_date else ""
                cells_lease_project_table[5].find('.//' + qn('w:t')).text = str(int(lease.vat))
                cells_lease_project_table[6].find('.//' + qn('w:t')).text = str(lease.vade)
                cells_lease_project_table[7].find('.//' + qn('w:t')).text = format_currency(lease.installment_amount)
                cells_lease_project_table[8].find('.//' + qn('w:t')).text = lease.currency.code if lease.currency.code != "TRY" else "TL"
                project_table._tbl.append(lease_row_project_table)

                #risk bilgileri tablosu___________________________________________________________________________________________
                risk_row_risk_table = deepcopy(template_lease_row_risk_table._tr)
                cells_risk_risk_table = risk_row_risk_table.findall(qn('w:tc'))
                cells_risk_risk_table[0].find('.//' + qn('w:t')).text = lease.contract.code if lease.contract else ""
                cells_risk_risk_table[1].find('.//' + qn('w:t')).text = datetime.today().strftime('%d.%m.%Y')
                cells_risk_risk_table[2].find('.//' + qn('w:t')).text = str(lease.overdue_days)
                cells_risk_risk_table[3].find('.//' + qn('w:t')).text = format_currency(next_installments['total_amount'] or Decimal('0.00'))
                cells_risk_risk_table[4].find('.//' + qn('w:t')).text = format_currency(lease.overdue_amount)
                cells_risk_risk_table[5].find('.//' + qn('w:t')).text = format_currency((next_installments['total_amount'] or Decimal('0.00')) + lease.overdue_amount)
                cells_risk_risk_table[6].find('.//' + qn('w:t')).text = format_currency(lease.paid_amount)
                risk_table._tbl.append(risk_row_risk_table)

            #toplam satırı ekleme
            #projeler tablosu___________________________________________________________________________________________
            project_total_amount = leases.aggregate(total_amount=Sum('installment_amount'))['total_amount'] or Decimal('0.00')
            project_total_row = deepcopy(template_total_row_project_table._tr)
            cells_total_project_table = project_total_row.findall(qn('w:tc'))
            cells_total_project_table[6].find('.//' + qn('w:t')).text = "Toplam"
            cells_total_project_table[7].find('.//' + qn('w:t')).text = format_currency(project_total_amount)
            cells_total_project_table[8].find('.//' + qn('w:t')).text = lease.currency.code if lease.currency.code != "TRY" else "TL"
            project_table._tbl.append(project_total_row)
            #risk bilgileri tablosu___________________________________________________________________________________________
            risk_total_row = deepcopy(template_total_row_risk_table._tr)
            cells_total_risk_table = risk_total_row.findall(qn('w:tc'))
            cells_total_risk_table[2].find('.//' + qn('w:t')).text = "Toplam"
            cells_total_risk_table[3].find('.//' + qn('w:t')).text = format_currency(risk_next_installments_total_amount)
            cells_total_risk_table[4].find('.//' + qn('w:t')).text = format_currency(risk_overdue_amount_total)
            cells_total_risk_table[5].find('.//' + qn('w:t')).text = format_currency(risk_total_risk_amount_total)
            cells_total_risk_table[6].find('.//' + qn('w:t')).text = format_currency(risk_paid_amount_total)
            risk_table._tbl.append(risk_total_row)

            #kesinti tablosu___________________________________________________________________________________________
            refund_installment_amount_row = deepcopy(template_installment_amount_row_refund_table._tr)
            cells_installment_amount_refund_table = refund_installment_amount_row.findall(qn('w:tc'))
            cells_installment_amount_refund_table[1].find('.//' + qn('w:t')).text = format_currency(refund_installment_amount_total)
            refund_table._tbl.append(refund_installment_amount_row)
            #++++++++++++
            refund_paid_amount_row = deepcopy(template_paid_amount_row_refund_table._tr)
            cells_paid_amount_refund_table = refund_paid_amount_row.findall(qn('w:tc'))
            cells_paid_amount_refund_table[1].find('.//' + qn('w:t')).text = format_currency(tahsil_edilen)
            refund_table._tbl.append(refund_paid_amount_row)
            #++++++++++++
            refund_warning_notice_row = deepcopy(template_warning_notice_row_refund_table._tr)
            cells_warning_notice_refund_table = refund_warning_notice_row.findall(qn('w:tc'))
            cells_warning_notice_refund_table[1].find('.//' + qn('w:t')).text = format_currency(ihtar_masrafi)
            refund_table._tbl.append(refund_warning_notice_row)
            #++++++++++++
            refund_tracking_row = deepcopy(template_tracking_row_refund_table._tr)
            cells_tracking_refund_table = refund_tracking_row.findall(qn('w:tc'))
            cells_tracking_refund_table[1].find('.//' + qn('w:t')).text = format_currency(takip_masrafi)
            refund_table._tbl.append(refund_tracking_row)
            #++++++++++++
            refund_payout_row = deepcopy(template_payout_row_refund_table._tr)
            cells_payout_refund_table = refund_payout_row.findall(qn('w:tc'))
            cells_payout_refund_table[1].find('.//' + qn('w:t')).text = format_currency(vazgecme_akcesi)
            refund_table._tbl.append(refund_payout_row)
            #++++++++++++
            refund_remaining_row = deepcopy(template_remaining_row_refund_table._tr)
            cells_remaining_refund_table = refund_remaining_row.findall(qn('w:tc'))
            cells_remaining_refund_table[1].find('.//' + qn('w:t')).text = format_currency(kalan)
            refund_table._tbl.append(refund_remaining_row)

            #şablon satırı silmek
            project_table._tbl.remove(template_lease_row_project_table._tr)
            project_table._tbl.remove(template_total_row_project_table._tr)
            risk_table._tbl.remove(template_lease_row_risk_table._tr)
            risk_table._tbl.remove(template_total_row_risk_table._tr)
            refund_table._tbl.remove(template_installment_amount_row_refund_table._tr)
            refund_table._tbl.remove(template_paid_amount_row_refund_table._tr)
            refund_table._tbl.remove(template_warning_notice_row_refund_table._tr)
            refund_table._tbl.remove(template_tracking_row_refund_table._tr)
            refund_table._tbl.remove(template_payout_row_refund_table._tr)
            refund_table._tbl.remove(template_remaining_row_refund_table._tr)
            document.save(files_path)

            if files_path:
                return FileResponse(open(files_path, 'rb'), as_attachment=True)
        
        return JsonResponse({'message': 'Dosya bulunamadı!','status':'error'}, status=400)
    
class GetCommitteeFormView(LoginRequiredMixin,View):
    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)
        uuid = data.get('uuid')

        partner = Partner.objects.select_related().filter(uuid = uuid).first()

        file_name = f"{partner.name.lower().replace(' ', '_')}-{str(partner.crm_code)}-komite-formu"
        file_path = os.path.join(settings.BASE_DIR, "media", "docs", str(self.request.user.user_companies.filter(is_active = True).first().company.uuid), "risk", "to_terminated_risk_partners", "documents",f"{file_name}.docx")
        
        if not os.path.exists(file_path):
            return JsonResponse({'message': 'Dosya bulunamadı!','status':'error'}, status=404)
        
        return FileResponse(open(file_path, 'rb'))
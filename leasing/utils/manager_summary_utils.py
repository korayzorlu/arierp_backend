from django.shortcuts import render
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import Sum,Count,Case,When,Value,BooleanField,Max
from django.views import View
from django.http import JsonResponse, FileResponse, HttpResponse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils.crypto import get_random_string
from django.conf import settings
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync

from utils.mixins import CompanyOwnershipRequiredMixin

from leasing.models import *
from common.utils.websocket_utils import send_alert
from common.utils.common_utils import parse_amount
from partners.models import Partner
from partners.utils.partner_utils import send_warning_email_for_ignored_partners
from contracts.models import Contract
from companies.models import UserCompany
from common.models import ExportProcess
from projects.models import Project

import os
import json
import pandas as pd
from decimal import Decimal
from datetime import datetime
from collections import Counter
from docxtpl import DocxTemplate
import io, zipfile
from copy import deepcopy
from docx.oxml.ns import qn

def format_currency(value):
    return "{:,.2f}".format(value).replace(",", "X").replace(".", ",").replace("X", ".")

def export_manager_summary(self):
    partner = Partner.objects.select_related().filter(uuid = self.params.get('partner')).first()
    #leases = Lease.objects.select_related().filter(uuid__in = self.params.get('uuids', []))
    leases = Lease.objects.select_related().filter(contract__partner = partner, is_last_project = True, is_last_project_arinet = True)

    self.process.status = "in_progress"
    self.process.items_count = len(leases)
    self.process.save()
    
    if partner:
        # word işlemleri
        file_name = f"{partner.name.lower().replace(' ', '_')}-{str(partner.crm_code)}-yonetici-ozeti"
        doc = DocxTemplate(f"files/yonetici-ozeti.docx")

        #hesaplamalar
        risk_next_installments_total_amount = Decimal('0.00')
        risk_overdue_amount_total = Decimal('0.00')
        risk_total_risk_amount_total = Decimal('0.00')
        risk_paid_amount_total = Decimal('0.00')
        refund_installment_amount_total = Decimal('0.00')
        muhtelif_transactions_total_amount = Decimal('0.00')
        teblig_tarihi_list = []
        for lease in leases:
            next_installments = Installment.objects.select_related("lease").filter(lease = lease, type__in = ['1'], payment_date__gte = datetime.today()).only("amount").aggregate(total_amount=Sum('amount'))
            borc_muhtelif_transactions = lease.lease_trade_transactions.filter(amount_type = '1', posting_group_name='Muhtelif Masraf').aggregate(total_amount=Sum('amount'))
            alacak_muhtelif_transactions = lease.lease_trade_transactions.filter(amount_type = '0', posting_group_name='Muhtelif Masraf').aggregate(total_amount=Sum('amount'))
            muhtelif_transactions_total_amount += (borc_muhtelif_transactions['total_amount'] or Decimal('0.00')) - (alacak_muhtelif_transactions['total_amount'] or Decimal('0.00'))
            
            risk_next_installments_total_amount += next_installments['total_amount'] or Decimal('0.00')
            risk_overdue_amount_total += lease.overdue_amount
            risk_total_risk_amount_total += (next_installments['total_amount'] or Decimal('0.00')) + lease.overdue_amount
            risk_paid_amount_total += lease.paid_amount
            refund_installment_amount_total += lease.installment_amount

        tahsil_edilen = risk_paid_amount_total
        ihtar_masrafi = Decimal('3200.00') if not partner.kep else Decimal('1000.00')*Decimal(str(leases.count()))
        takip_masrafi = risk_overdue_amount_total*Decimal('0.1') if risk_overdue_amount_total > muhtelif_transactions_total_amount else muhtelif_transactions_total_amount
        vazgecme_akcesi = refund_installment_amount_total*Decimal('0.1')
        kalan = risk_paid_amount_total - ihtar_masrafi - takip_masrafi - vazgecme_akcesi

        context = {
            "isim": partner.name,
            "adres": partner.address,
            "tel": partner.phone_number,
            "tarih": datetime.today().strftime('%d.%m.%Y'),
            "kalan": format_currency(kalan),
            "pb": leases[0].currency.code if leases and leases[0].currency.code != "TRY" else "TL",
        }

        doc.render(context)

        base_path = os.path.join(os.getcwd(), "media", "docs", str(self.user.user_companies.filter(is_active=True).first().company.uuid), "leasing", "manager_summary", "documents")
        if not os.path.exists(base_path):
            os.makedirs(base_path)

        files_path = os.path.join(base_path, f"{file_name}.docx")
        doc.save(files_path)

        # tabloya satır ekleme
        from docx import Document
        document = Document(files_path)

        project_table = document.tables[2]
        template_lease_row_project_table = project_table.rows[2]
        template_total_row_project_table = project_table.rows[3]

        risk_table = document.tables[3]
        template_lease_row_risk_table = risk_table.rows[2]
        template_total_row_risk_table = risk_table.rows[3]

        for lease in leases:
            #projeler tablosu___________________________________________________________________________________________
            lease_row_project_table = deepcopy(template_lease_row_project_table._tr)
            cells_lease_project_table = lease_row_project_table.findall(qn('w:tc'))
            cells_lease_project_table[0].find('.//' + qn('w:t')).text = lease.contract.vendor.name if lease.contract and lease.contract.vendor else ""
            cells_lease_project_table[1].find('.//' + qn('w:t')).text = lease.item.stock_name if lease.item else ""
            cells_lease_project_table[2].find('.//' + qn('w:t')).text = lease.contract.code if lease.contract else ""
            cells_lease_project_table[3].find('.//' + qn('w:t')).text = lease.get_lease_status_display()
            cells_lease_project_table[4].find('.//' + qn('w:t')).text = lease.signature_date.strftime('%d.%m.%Y') if lease.signature_date else ""
            cells_lease_project_table[5].find('.//' + qn('w:t')).text = str(int(lease.vat))
            cells_lease_project_table[6].find('.//' + qn('w:t')).text = str(lease.vade)
            cells_lease_project_table[7].find('.//' + qn('w:t')).text = format_currency(lease.installment_amount)
            cells_lease_project_table[8].find('.//' + qn('w:t')).text = lease.currency.code if lease.currency.code != "TRY" else "TL"
            project_table._tbl.append(lease_row_project_table)

            #risk bilgileri tablosu___________________________________________________________________________________________
            risk_row_risk_table = deepcopy(template_lease_row_risk_table._tr)
            cells_risk_risk_table = risk_row_risk_table.findall(qn('w:tc'))
            cells_risk_risk_table[0].find('.//' + qn('w:t')).text = lease.contract.code if lease.contract else ""
            cells_risk_risk_table[1].find('.//' + qn('w:t')).text = datetime.today().strftime('%d.%m.%Y')
            cells_risk_risk_table[2].find('.//' + qn('w:t')).text = str(lease.overdue_days)
            cells_risk_risk_table[3].find('.//' + qn('w:t')).text = format_currency(next_installments['total_amount'] or Decimal('0.00'))
            cells_risk_risk_table[4].find('.//' + qn('w:t')).text = format_currency(lease.overdue_amount)
            cells_risk_risk_table[5].find('.//' + qn('w:t')).text = format_currency((next_installments['total_amount'] or Decimal('0.00')) + lease.overdue_amount)
            cells_risk_risk_table[6].find('.//' + qn('w:t')).text = format_currency(lease.paid_amount)
            risk_table._tbl.append(risk_row_risk_table)

        #toplam satırı ekleme
        #projeler tablosu___________________________________________________________________________________________
        project_total_amount = leases.aggregate(total_amount=Sum('installment_amount'))['total_amount'] or Decimal('0.00')
        project_total_row = deepcopy(template_total_row_project_table._tr)
        cells_total_project_table = project_total_row.findall(qn('w:tc'))
        cells_total_project_table[6].find('.//' + qn('w:t')).text = "Toplam"
        cells_total_project_table[7].find('.//' + qn('w:t')).text = format_currency(project_total_amount)
        cells_total_project_table[8].find('.//' + qn('w:t')).text = leases[0].currency.code if leases and leases[0].currency.code != "TRY" else "TL"
        project_table._tbl.append(project_total_row)
        #risk bilgileri tablosu___________________________________________________________________________________________
        risk_total_row = deepcopy(template_total_row_risk_table._tr)
        cells_total_risk_table = risk_total_row.findall(qn('w:tc'))
        cells_total_risk_table[2].find('.//' + qn('w:t')).text = "Toplam"
        cells_total_risk_table[3].find('.//' + qn('w:t')).text = format_currency(risk_next_installments_total_amount)
        cells_total_risk_table[4].find('.//' + qn('w:t')).text = format_currency(risk_overdue_amount_total)
        cells_total_risk_table[5].find('.//' + qn('w:t')).text = format_currency(risk_total_risk_amount_total)
        cells_total_risk_table[6].find('.//' + qn('w:t')).text = format_currency(risk_paid_amount_total)
        risk_table._tbl.append(risk_total_row)

        #şablon satırı silmek
        project_table._tbl.remove(template_lease_row_project_table._tr)
        project_table._tbl.remove(template_total_row_project_table._tr)
        risk_table._tbl.remove(template_lease_row_risk_table._tr)
        risk_table._tbl.remove(template_total_row_risk_table._tr)
        document.save(files_path)
        
    self.process.progress = 100
    #self.process.status = "completed"
    self.process.save()